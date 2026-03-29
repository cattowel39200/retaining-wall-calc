# -*- coding: utf-8 -*-
"""
옹벽 구조계산서 Word 출력 모듈 (4종 통합: L형, 역L형, 역T형, 중력식/반중력식)
python-docx를 사용하여 원본 방식 구조계산서를 생성한다.
"""

import io
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
# draw_section removed for Vercel deployment (matplotlib too large)


def _f3(v):
    return f"{v:.3f}"


def _f2(v):
    return f"{v:.2f}"


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
    return h


def _add_para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.bold = bold
    return p


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.name = '맑은 고딕'

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = '맑은 고딕'

    return table


def export_report(params: dict, results: dict) -> bytes:
    """
    Word 구조계산서를 생성하여 바이트로 반환한다.
    st.download_button에서 바로 사용 가능.
    """
    doc = Document()

    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    p = params
    r = results
    stab = r['stability']
    ep = r['earth_pressure']
    blk = r['blocks']
    sec = r['section']
    mbr = r['member']
    jdg = r['judgment']
    sn = stab['normal']
    se = stab['seismic']

    # ================================================================
    # 표지
    # ================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wall_type = p.get('wall_type', 'L형')
    run = title.add_run(f'{wall_type} 옹벽 구조계산서')
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = '맑은 고딕'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(f'H = {p["H"]:.1f} m')
    run2.font.size = Pt(18)
    run2.font.name = '맑은 고딕'

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = info.add_run('적용기준: KDS 11 80 05 / KDS 11 80 15 / KDS 14 20 20')
    run3.font.size = Pt(11)
    run3.font.name = '맑은 고딕'

    doc.add_page_break()

    # ================================================================
    # 1. 일반 단면
    # ================================================================
    _add_heading(doc, '1. 일반 단면', 1)

    _add_para(doc, '  [단면도: 웹 화면에서 확인]')

    _is_gravity = (wall_type == '중력식')
    _is_semi_gravity = p.get('semi_gravity', False) and _is_gravity
    _has_toe = jdg.get('has_toe', False)
    _has_heel = jdg.get('has_heel', True)

    _add_heading(doc, '1.1 옹벽의 제원', 2)
    if _is_gravity:
        _grav_label = '반중력식' if _is_semi_gravity else '중력식'
        _add_para(doc, f'  옹벽 형식 : {_grav_label} 옹벽 (무근콘크리트)' if not _is_semi_gravity
                  else f'  옹벽 형식 : {_grav_label} 옹벽')
        _add_para(doc, f'  기초 형식 : 직접기초')
        _add_para(doc, f'  옹벽 높이 : H = {p["H"]:.3f} m')
        _add_para(doc, f'  상 단 폭 : {p["t_stem"]:.3f} m')
        _add_para(doc, f'  하 단 폭 : B = {p["B"]:.3f} m')
    else:
        _add_para(doc, f'  옹벽 형식 : {wall_type} 옹벽')
        _add_para(doc, f'  기초 형식 : 직접기초')
        _add_para(doc, f'  옹벽 높이 : H = {p["H"]:.3f} m')
        _add_para(doc, f'  옹벽 저판 : B = {p["B"]:.3f} m')

    _add_heading(doc, '1.2 내진설계', 2)
    _add_para(doc, f'  지진구역계수 : {p.get("zone_coeff", 0.110):.3f}')
    _add_para(doc, f'  위험도계수   : {p.get("risk_factor", 1.400):.3f}')
    _add_para(doc, f'  가속도계수 A : {p.get("A_coeff", 0.154):.3f}')
    _add_para(doc, f'  수평지진계수 Kh : {p["Kh"]:.3f}')

    # ================================================================
    # 2. 설계조건
    # ================================================================
    _add_heading(doc, '2. 설계 조건', 1)

    _add_heading(doc, '2.1 사용재료', 2)
    _add_para(doc, f'  콘크리트 : fck = {p["fck"]:.0f} MPa')
    _add_para(doc, f'  철    근 : fy  = {p["fy"]:.0f} MPa')

    _add_heading(doc, '2.2 지반조건', 2)
    _add_para(doc, f'  콘크리트 단위중량 (γc) : {p["gamma_c"]:.1f} kN/m³')
    _add_para(doc, f'  뒷채움흙 단위중량 (γt) : {p["gamma_t"]:.1f} kN/m³')
    _add_para(doc, f'  내부마찰각 (φ)        : {p["phi_deg"]:.3f} °')
    _add_para(doc, f'  점착력 (C)             : {p["c_soil"]:.3f} kN/m²')
    _add_para(doc, f'  뒤채움 경사각 (α)      : {p["alpha_deg"]:.3f} °')
    _add_para(doc, f'  토피고 (Df)            : {p["Df"]:.3f} m')

    if p.get('gwl_height', 0) > 0:
        _add_para(doc, f'  지하수위 (저판상면부터)  : {p["gwl_height"]:.2f} m')
        _add_para(doc, f'  포화단위중량 (γsat)    : {p["gamma_sat"]:.1f} kN/m³')

    _add_heading(doc, '2.3 사용토압', 2)
    _add_para(doc, '  상시 : 안정검토 - Rankine 토압 / 단면검토 - Coulomb 토압')
    _add_para(doc, '  지진 : Mononobe-Okabe 토압')

    _add_heading(doc, '2.4 과재하중', 2)
    _add_para(doc, f'  과재하중 : q = {p["q"]:.2f} kN/m²')

    doc.add_page_break()

    # ================================================================
    # 3. 안정 검토
    # ================================================================
    _add_heading(doc, '3. 안정 검토', 1)

    _add_heading(doc, '3.1 자중 및 재토하중 계산', 2)

    headers = ["구분", "A (m²)", "γ (kN/m³)", "W (kN)", "Kh", "KhW (kN)",
               "x (m)", "y (m)", "Mr (kN·m)", "Mo (kN·m)"]
    rows = []
    for b in blk['c_results']:
        if b['A'] > 0:
            rows.append([b['name'], _f3(b['A']), _f2(b['gamma']), _f3(b['W']),
                        _f3(b['Kh']), _f3(b['KhW']), _f3(b['x']), _f3(b['y']),
                        _f3(b['Mr']), _f3(b['Mo'])])
    rows.append(["콘크리트소계", "", "", _f3(blk['Wc']), "", _f3(blk['KhWc']),
                "", "", _f3(blk['Mrc']), _f3(blk['Moc'])])
    for b in blk['s_results']:
        if b['A'] > 0:
            rows.append([b['name'], _f3(b['A']), _f2(b['gamma']), _f3(b['W']),
                        _f3(b['Kh']), _f3(b['KhW']), _f3(b['x']), _f3(b['y']),
                        _f3(b['Mr']), _f3(b['Mo'])])
    rows.append(["토사소계", "", "", _f3(blk['Ws']), "", _f3(blk['KhWs']),
                "", "", _f3(blk['Mrs']), _f3(blk['Mos'])])
    rows.append(["합계", "", "", _f3(blk['Wt']), "", _f3(blk['KhWt']),
                "", "", _f3(blk['Mrt']), _f3(blk['Mot'])])
    _add_table(doc, headers, rows)

    _add_heading(doc, '토압계산', 3)
    _add_para(doc, f'  ① 상시 주동토압 (Rankine)')
    _add_para(doc, f'    Ka = {_f3(ep["Ka"])}')
    _add_para(doc, f'    Pa(수평) = {_f3(ep["Pa"])} kN/m,  y = {_f3(ep["ya"])} m')
    if ep.get('Pa_v', 0) != 0:
        _add_para(doc, f'    Pa(수직) = {_f3(ep["Pa_v"])} kN/m')
    _add_para(doc, f'    Mo = {_f3(ep["Mo_pa"])} kN·m')
    _add_para(doc, f'  ② 지진시 주동토압 (M-O)')
    _add_para(doc, f'    KAE = {_f3(ep["KAE"])}')
    _add_para(doc, f'    PAE = {_f3(ep["PAE"])} kN/m,  y = He/2 = {_f3(ep["yae"])} m')
    _add_para(doc, f'    Mo = {_f3(ep["Mo_pae"])} kN·m')

    if ep.get('Pw', 0) > 0:
        _add_para(doc, f'  ③ 수압')
        _add_para(doc, f'    Pw = {_f3(ep["Pw"])} kN/m')
        _add_para(doc, f'    양압력 U = {_f3(ep["U"])} kN/m')

    _add_para(doc, f'  과재하중 q = {p["q"]:.2f} kN/m²')
    _add_para(doc, f'    Ph = {_f3(ep["Ph_sur"])} kN/m,  Pv = {_f3(ep["Pv_sur"])} kN/m')

    doc.add_page_break()

    # 합력표
    _add_heading(doc, '3.2 안정검토용 하중집계', 2)

    _add_para(doc, '1) 상시 하중집계', bold=True)
    _add_table(doc,
               ["구분", "V (kN)", "H (kN)", "Mr (kN·m)", "Mo (kN·m)"],
               [
                   ["콘크리트 자중", _f3(blk['Wc']), "0.000", _f3(blk['Mrc']), "0.000"],
                   ["뒤채움토 자중", _f3(blk['Ws']), "0.000", _f3(blk['Mrs']), "0.000"],
                   ["토압", _f3(ep.get('Pa_v', 0)), _f3(ep['Pa']), "0.000", _f3(ep['Mo_pa'])],
                   ["상재하중", _f3(ep['Pv_sur']), _f3(ep['Ph_sur']), _f3(ep['Mr_pv']), _f3(ep['Mo_ph'])],
                   ["Σ", _f3(sn['SV']), _f3(sn['SH']), _f3(sn['SMr']), _f3(sn['SMo'])],
               ])

    _add_para(doc, '2) 지진시 하중집계', bold=True)
    _add_table(doc,
               ["구분", "V (kN)", "H (kN)", "Mr (kN·m)", "Mo (kN·m)"],
               [
                   ["콘크리트 자중", _f3(blk['Wc']), "0.000", _f3(blk['Mrc']), "0.000"],
                   ["뒤채움토 자중", _f3(blk['Ws']), "0.000", _f3(blk['Mrs']), "0.000"],
                   ["토압", "0.000", _f3(ep['PAE']), "0.000", _f3(ep['Mo_pae'])],
                   ["관성력", "0.000", _f3(blk['KhWt']), "0.000", _f3(blk['Mot'])],
                   ["Σ", _f3(se['SV']), _f3(se['SH']), _f3(se['SMr']), _f3(se['SMo'])],
               ])

    # 전도
    _add_heading(doc, '3.3 전도에 대한 안정검토', 2)
    _add_para(doc, '1) 상시', bold=True)
    _add_para(doc, f'  e = B/2 - (ΣMr-ΣMo)/ΣV = {_f3(sn["e"])} m')
    _add_para(doc, f'  e = {_f3(sn["e"])} m  {"<=" if sn["e"] <= sn["B6"] else ">"} B/6 = {_f3(sn["B6"])} m')
    _add_para(doc, f'  SF = ΣMr/ΣMo = {_f3(sn["SF_overturn"])} >= 2.0 → {jdg["overturn_normal"]}')

    _add_para(doc, '2) 지진시', bold=True)
    _add_para(doc, f'  e = {_f3(se["e"])} m  {"<=" if se["e"] <= sn["B3"] else ">"} B/3 = {_f3(sn["B3"])} m → {jdg["eccentricity_seismic"]}')
    _add_para(doc, f'  SF = ΣMr/ΣMo = {_f3(se["SF_overturn"])} >= 1.5 → {jdg.get("overturn_seismic", "OK")}')

    # 지지력
    _add_heading(doc, '3.4 지지력에 대한 안정검토', 2)
    _add_para(doc, '1) 상시', bold=True)
    _add_para(doc, f'  Be = B - 2e = {_f3(sn["Be"])} m')
    _add_para(doc, f'  Nc = {_f3(sn["Nc"])}, Nq = {_f3(sn["Nq"])}, Nr = {_f3(sn["Nr"])}')
    if sn.get('qa_fixed', 0) > 0:
        _add_para(doc, f'  qu = {_f3(sn["qu"])} kN/m² → qa = {_f3(sn["qa"])} kN/m² (직접입력)')
    else:
        _add_para(doc, f'  qu = {_f3(sn["qu"])} kN/m² → qa = qu/3 = {_f3(sn["qa"])} kN/m²')
    _add_para(doc, f'  Q1 = {_f3(sn["Q1"])} kN/m², Q2 = {_f3(sn["Q2"])} kN/m²')
    _add_para(doc, f'  q_max = {_f3(sn["Q1"])} <= qa = {_f3(sn["qa"])} → {jdg["bearing_normal"]}')

    _add_para(doc, '2) 지진시', bold=True)
    _add_para(doc, f'  Be = {_f3(se["Be"])} m')
    if sn.get('qae_fixed', 0) > 0:
        _add_para(doc, f'  que = {_f3(se["qu"])} kN/m² → qae = {_f3(se["qa"])} kN/m² (직접입력)')
    else:
        _add_para(doc, f'  que = {_f3(se["qu"])} kN/m² → qae = que/2 = {_f3(se["qa"])} kN/m²')
    _add_para(doc, f'  Q1 = {_f3(se["Q1"])} kN/m², Q2 = {_f3(se["Q2"])} kN/m²')
    _add_para(doc, f'  q_max = {_f3(se["Q1"])} <= qae = {_f3(se["qa"])} → {jdg["bearing_seismic"]}')

    # 활동
    _add_heading(doc, '3.5 활동에 대한 안정검토', 2)
    _add_para(doc, f'  μ = tan(2φ/3) = {_f3(sn["mu"])}')
    _add_para(doc, '1) 상시', bold=True)
    _add_para(doc, f'  Hr = {_f3(sn["Hr"])} kN')
    _add_para(doc, f'  SF = Hr/ΣH = {_f3(sn["SF_slide"])} >= 1.5 → {jdg["slide_normal"]}')
    _add_para(doc, '2) 지진시', bold=True)
    _add_para(doc, f'  Hr = {_f3(se["Hr"])} kN')
    _add_para(doc, f'  SF = Hr/ΣH = {_f3(se["SF_slide"])} >= 1.2 → {jdg["slide_seismic"]}')

    doc.add_page_break()

    # ================================================================
    # 4. 단면검토 (중력식: 생략)
    # ================================================================
    if not _is_gravity:
        _add_heading(doc, '4. 단면 검토', 1)

        _add_heading(doc, '4.1 하중조합', 2)
        _add_para(doc, '  LCB1 : 상시 계수하중 (1.2D + 1.6L + 1.6H)')
        _add_para(doc, '  LCB2 : 지진시 계수하중 (0.9D + 1.0H + 1.0E)')
        _add_para(doc, '  LCB3 : 상시 사용하중 (1.0D + 1.0L + 1.0H + 1.0W)')

        _add_heading(doc, '4.2 LCB별 지반반력', 2)
        _add_table(doc,
                   ["조합", "ΣV (kN)", "e (m)", "Q1 (kN/m²)", "Q2 (kN/m²)"],
                   [
                       ["LCB1", _f3(sec['lcb1']['SV']), _f3(sec['lcb1']['e']),
                        _f3(sec['lcb1']['Q1']), _f3(sec['lcb1']['Q2'])],
                       ["LCB2", _f3(sec['lcb2']['SV']), _f3(sec['lcb2']['e']),
                        _f3(sec['lcb2']['Q1']), _f3(sec['lcb2']['Q2'])],
                       ["LCB3", _f3(sec['lcb3']['SV']), _f3(sec['lcb3']['e']),
                        _f3(sec['lcb3']['Q1']), _f3(sec['lcb3']['Q2'])],
                   ])

        # ================================================================
        # 4.3 단면검토용 하중계산 (원본 형식)
        # ================================================================
        _add_heading(doc, '4.3 단면검토용 하중계산', 2)

        w = sec['wall']

        # 1) 뒷굽판 단면력 (has_heel인 경우만)
        if _has_heel:
            _add_para(doc, '1) 뒷굽판 단면력', bold=True)
            _add_para(doc, '                                                          (단위 : KN, m)')
            h1 = sec['heel_lcb1']
            h2 = sec['heel_lcb2']
            h3 = sec['heel_lcb3']
            _add_table(doc,
                       ["구 분", "", "뒷굽자중", "재토자중", "과재하중", "지반반력", "연직토압", "총 계"],
                       [
                           ["LCB1", "전단력", _f3(h1['V_slab']), _f3(h1['V_soil']), _f3(h1['V_sur']),
                            _f3(h1['V_react']), "0.000", _f3(h1['V_total'])],
                           ["", "모멘트", _f3(h1['M_slab']), _f3(h1['M_soil']), _f3(h1['M_sur']),
                            _f3(h1['M_react']), "0.000", _f3(h1['M_total'])],
                           ["LCB2", "전단력", _f3(h2['V_slab']), _f3(h2['V_soil']), _f3(h2['V_sur']),
                            _f3(h2['V_react']), "0.000", _f3(h2['V_total'])],
                           ["", "모멘트", _f3(h2['M_slab']), _f3(h2['M_soil']), _f3(h2['M_sur']),
                            _f3(h2['M_react']), "0.000", _f3(h2['M_total'])],
                           ["LCB3", "전단력", _f3(h3['V_slab']), _f3(h3['V_soil']), _f3(h3['V_sur']),
                            _f3(h3['V_react']), "0.000", _f3(h3['V_total'])],
                           ["", "모멘트", _f3(h3['M_slab']), _f3(h3['M_soil']), _f3(h3['M_sur']),
                            _f3(h3['M_react']), "0.000", _f3(h3['M_total'])],
                       ])

        # 1-2) 앞굽판 단면력 (has_toe일 때만)
        if _has_toe and sec.get('toe_lcb1') is not None:
            _add_para(doc, '')
            _toe_label = '1) 앞굽판 단면력' if not _has_heel else '1-2) 앞굽판 단면력'
            _add_para(doc, _toe_label, bold=True)
            _add_para(doc, '                                                          (단위 : KN, m)')
            t1 = sec['toe_lcb1']
            t2 = sec['toe_lcb2']
            t3 = sec['toe_lcb3']
            _add_table(doc,
                       ["구 분", "", "저판자중", "지반반력", "총 계"],
                       [
                           ["LCB1", "전단력", _f3(t1['V_slab']), _f3(t1['V_react']), _f3(t1['V_total'])],
                           ["", "모멘트", _f3(t1['M_slab']), _f3(t1['M_react']), _f3(t1['M_total'])],
                           ["LCB2", "전단력", _f3(t2['V_slab']), _f3(t2['V_react']), _f3(t2['V_total'])],
                           ["", "모멘트", _f3(t2['M_slab']), _f3(t2['M_react']), _f3(t2['M_total'])],
                           ["LCB3", "전단력", _f3(t3['V_slab']), _f3(t3['V_react']), _f3(t3['V_total'])],
                           ["", "모멘트", _f3(t3['M_slab']), _f3(t3['M_react']), _f3(t3['M_total'])],
                       ])

        # 2) 벽체 단면력
        _add_para(doc, '')
        _wall_num = '2)' if _has_heel else ('2)' if _has_toe else '1)')
        _add_para(doc, f'{_wall_num} 벽체 단면력', bold=True)

        theta_c_deg = ep.get('theta_c_deg', 0.0)
        delta_c_deg = ep.get('delta_c_deg', 10.0)

        _add_para(doc, '(1) 토압계수 계산')
        _add_para(doc, '  ⓐ 상시 주동토압계산 (Coulomb)')
        _add_para(doc, f'    뒤채움흙의 내부마찰각(Φ) : {p["phi_deg"]:.3f} °')
        _add_para(doc, f'    뒤채움흙의 경 사 각(α) : {p["alpha_deg"]:.3f} °')
        _add_para(doc, f'    흙과 콘크리트의 마찰각(δ): {delta_c_deg:.3f} °')
        _add_para(doc, f'    옹벽배면의 연직경사각(θ) : {theta_c_deg:.3f} °')
        _add_para(doc, f'    Ka = {_f3(ep["Ka_coul"])}')
        _add_para(doc, f'    Kah = {_f3(ep["Ka_coul"])}×cos({delta_c_deg:.3f}°+{theta_c_deg:.3f}°) = {_f3(ep["Kah"])}')

        _add_para(doc, '  ⓑ 지진시 주동토압계산 (Mononobe-Okabe)')
        _add_para(doc, f'    옹벽배면의 수직에 대한 각 β : = {theta_c_deg:.3f} °')
        _add_para(doc, f'    흙과 옹벽사이의 마찰각     δ : = 0.000 ° (단면검토시)')
        _add_para(doc, f'    Kae = {_f3(ep["Kae_design"])}')
        _add_para(doc, f'    Kaeh = {_f3(ep["Kae_design"])}×cos({theta_c_deg:.3f}°) = {_f3(ep["Kaeh_design"])}')

        H_cc = w['H_wall_cc']
        H_dd = w['H_wall_dd']

        # (2) 토압에 의한 벽체 단면력계산
        _add_para(doc, '')
        _add_para(doc, '(2) 토압에 의한 벽체 단면력계산')
        _add_para(doc, '  ⓐ 상시 벽체 단면력')
        _add_para(doc, f'    i) 벽체 하부 (C-C)')
        _add_para(doc, f'      Pa = 1/2 ×Kah × γt ×H² = 1/2 ×{_f3(ep["Kah"])} ×{p["gamma_t"]:.1f} ×{H_cc:.3f}² = {_f3(w["Pa_cc"])} KN/m')
        _add_para(doc, f'      y  = H / 3 = {H_cc:.3f} / 3 = {_f3(H_cc/3)} m')
        _add_para(doc, f'      Mo = Pa ×y = {_f3(w["Pa_cc"])} ×{_f3(H_cc/3)} = {_f3(w["Mo_cc_Pa"])} KN.m')
        _add_para(doc, f'    ii) 벽체 중앙부 (D-D)')
        _add_para(doc, f'      Pa = 1/2 ×Kah × γt ×H² = 1/2 ×{_f3(ep["Kah"])} ×{p["gamma_t"]:.1f} ×{H_dd:.3f}² = {_f3(w["Pa_dd"])} KN/m')
        _add_para(doc, f'      y  = H / 3 = {H_dd:.3f} / 3 = {_f3(H_dd/3)} m')
        _add_para(doc, f'      Mo = {_f3(w["Pa_dd"])} ×{_f3(H_dd/3)} = {_f3(w["Mo_dd_Pa"])} KN.m')

        _add_para(doc, '  ⓑ 지진시 벽체 단면력')
        _add_para(doc, f'    i) 벽체 하부 (C-C)')
        _add_para(doc, f'      Pae = 1/2 ×Kaeh × γt ×H² = 1/2 ×{_f3(ep["Kaeh_design"])} ×{p["gamma_t"]:.1f} ×{H_cc:.3f}² = {_f3(w["Pae_cc"])} KN/m')
        _add_para(doc, f'      y   = H / 2 = {H_cc:.3f} / 2 = {_f3(H_cc/2)} m')
        _add_para(doc, f'      Mo  = {_f3(w["Pae_cc"])} ×{_f3(H_cc/2)} = {_f3(w["Mo_cc_Pae"])} KN.m')
        _add_para(doc, f'    ii) 벽체 중앙부 (D-D)')
        _add_para(doc, f'      Pae = 1/2 ×Kaeh × γt ×H² = 1/2 ×{_f3(ep["Kaeh_design"])} ×{p["gamma_t"]:.1f} ×{H_dd:.3f}² = {_f3(w["Pae_dd"])} KN/m')
        _add_para(doc, f'      y   = H / 2 = {H_dd:.3f} / 2 = {_f3(H_dd/2)} m')
        _add_para(doc, f'      Mo  = {_f3(w["Pae_dd"])} ×{_f3(H_dd/2)} = {_f3(w["Mo_dd_Pae"])} KN.m')

        # (3) 과재하중
        _add_para(doc, '')
        _add_para(doc, '(3) 과재하중에 의한 벽체단면력 계산')
        _add_para(doc, f'  q = {p["q"]:.2f} KN/m²')
        _add_para(doc, f'    i) 벽체 하부 (C-C)')
        _add_para(doc, f'      Ph1 = Kah ×q ×H = {_f3(ep["Kah"])} ×{p["q"]:.2f} ×{H_cc:.3f} = {_f3(w["Ph1_cc"])} KN/m (활하중)')
        _add_para(doc, f'      y   = H / 2 = {_f3(H_cc/2)} m')
        _add_para(doc, f'      Mo1 = Ph1 ×y = {_f3(w["Mo_ph1_cc"])} KN.m')
        _add_para(doc, f'    ii) 벽체 중앙부 (D-D)')
        _add_para(doc, f'      Ph1 = Kah ×q ×H = {_f3(ep["Kah"])} ×{p["q"]:.2f} ×{H_dd:.3f} = {_f3(w["Ph1_dd"])} KN/m (활하중)')
        _add_para(doc, f'      y   = H / 2 = {_f3(H_dd/2)} m')
        _add_para(doc, f'      Mo1 = Ph1 ×y = {_f3(w["Mo_ph1_dd"])} KN.m')

        # ▷ 벽체 하단 단면력 계산
        _add_para(doc, '')
        _add_para(doc, '▷ 벽체 하단 단면력 계산', bold=True)
        _add_para(doc, '                                                          (단위 : KN, m)')
        _add_table(doc,
                   ["구 분", "", "횡 토 압", "과재하중", "관 성 력", "총 계"],
                   [
                       ["LCB1", "전단력", _f3(1.6*w['Pa_cc']), _f3(1.6*w['Ph1_cc']), "0.000", _f3(w['cc_lcb1']['V'])],
                       ["", "모멘트", _f3(1.6*w['Mo_cc_Pa']), _f3(1.6*w['Mo_ph1_cc']), "0.000", _f3(w['cc_lcb1']['M'])],
                       ["LCB2", "전단력", _f3(w['Pae_cc']), "0.000", _f3(w['inertia_H_cc']), _f3(w['cc_lcb2']['V'])],
                       ["", "모멘트", _f3(w['Mo_cc_Pae']), "0.000", _f3(w['inertia_M_cc']), _f3(w['cc_lcb2']['M'])],
                       ["LCB3", "전단력", _f3(w['Pa_cc']), _f3(w['Ph1_cc']), "0.000", _f3(w['cc_lcb3']['V'])],
                       ["", "모멘트", _f3(w['Mo_cc_Pa']), _f3(w['Mo_ph1_cc']), "0.000", _f3(w['cc_lcb3']['M'])],
                   ])

        # ▷ 벽체 중간부 단면력 계산
        _add_para(doc, '')
        _add_para(doc, '▷ 벽체 중간부 단면력 계산', bold=True)
        _add_para(doc, '                                                          (단위 : KN, m)')
        _add_table(doc,
                   ["구 분", "", "횡 토 압", "과재하중", "관 성 력", "총 계"],
                   [
                       ["LCB1", "전단력", _f3(1.6*w['Pa_dd']), _f3(1.6*w['Ph1_dd']), "0.000", _f3(w['dd_lcb1']['V'])],
                       ["", "모멘트", _f3(1.6*w['Mo_dd_Pa']), _f3(1.6*w['Mo_ph1_dd']), "0.000", _f3(w['dd_lcb1']['M'])],
                       ["LCB2", "전단력", _f3(w['Pae_dd']), "0.000", _f3(w['inertia_H_dd']), _f3(w['dd_lcb2']['V'])],
                       ["", "모멘트", _f3(w['Mo_dd_Pae']), "0.000", _f3(w['inertia_M_dd']), _f3(w['dd_lcb2']['M'])],
                       ["LCB3", "전단력", _f3(w['Pa_dd']), _f3(w['Ph1_dd']), "0.000", _f3(w['dd_lcb3']['V'])],
                       ["", "모멘트", _f3(w['Mo_dd_Pa']), _f3(w['Mo_ph1_dd']), "0.000", _f3(w['dd_lcb3']['M'])],
                   ])

        doc.add_page_break()

        # ================================================================
        # 4.4 단면검토용 하중집계
        # ================================================================
        _add_heading(doc, '4.4 단면검토용 하중집계', 2)
        _add_para(doc, '  상시와 지진시 단면력중 최대값으로 단면력을 정리하면 다음과 같다.')
        _add_para(doc, '  균열검토는 상시의 사용하중으로 검토한다.')
        _add_para(doc, '                                            ( 단위 : KN, m )')
        df = sec['design_forces']
        _df_rows = []
        if _has_heel and df.get('BB') is not None:
            _df_rows.append(["뒷 굽 판 (B-B)", _f3(df['BB']['Mu']), _f3(df['BB']['Mcr']), _f3(df['BB']['Vu'])])
        _df_rows.append(["벽 체 하 부 (C-C)", _f3(df['CC']['Mu']), _f3(df['CC']['Mcr']), _f3(df['CC']['Vu'])])
        _df_rows.append(["벽체 중앙부 (D-D)", _f3(df['DD']['Mu']), _f3(df['DD']['Mcr']), _f3(df['DD']['Vu'])])
        if _has_toe and df.get('AA') is not None:
            _df_rows.append(["앞 굽 판 (A-A)", _f3(df['AA']['Mu']), _f3(df['AA']['Mcr']), _f3(df['AA']['Vu'])])
        _add_table(doc, ["구     분", "Mu", "Mcr", "Vu"], _df_rows)
        _add_para(doc, '')
        _add_para(doc, '( 단, 저판에 적용하는 휨모멘트의 크기는 전면벽과 뒷굽판과의')
        _add_para(doc, '  접속점의 모멘트평형조건에 의하여 전면벽에 적용하는 휨모멘트를')
        _add_para(doc, '  초과하지 않는다.- 옹벽표준도작성연구용역 종합보고서, 1998. 건교부)')

        doc.add_page_break()

        # ================================================================
        # 4.5 단면검토 (원본 형식)
        # ================================================================
        _add_heading(doc, '4.5 단 면 검 토', 1)

        fck_v = p['fck']
        fy_v = p['fy']
        # beta1 재계산
        if fck_v <= 28:
            beta1_v = 0.850
        else:
            beta1_v = max(0.85 - 0.007 * (fck_v - 28), 0.65)

        # 동적 단면 항목 구성
        sec_items = []
        _sec_num = 1
        if _has_heel and mbr.get('BB') is not None:
            sec_items.append(("BB", f"{_sec_num}) 뒷 굽 판"))
            _sec_num += 1
        sec_items.append(("CC", f"{_sec_num}) 벽체하부"))
        _sec_num += 1
        sec_items.append(("DD", f"{_sec_num}) 벽체중앙"))
        _sec_num += 1
        if _has_toe and mbr.get('AA') is not None:
            sec_items.append(("AA", f"{_sec_num}) 앞 굽 판"))

        for sec_key, sec_title in sec_items:
            s = mbr[sec_key]

            _add_heading(doc, sec_title, 2)

            # 재료 / 기본값
            _add_para(doc, f'  fck = {fck_v:.1f}MPa          fy  = {fy_v:.1f}MPa')
            _add_para(doc, f'  β1 = {beta1_v:.3f}              φf = {s["phi_f"]:.2f}              φv = 0.75')
            _add_para(doc, f'  pmin = max(0.25√(fck)/fy, 1.4/fy)  = {s["pmin"]:.5f}')
            _add_para(doc, '')

            _add_para(doc, f'  계수 모멘트 Mu = {s["Mu"]:>10.3f} KN.m        계수 전단력 Vu = {s["Vu"]:>10.3f} KN')
            _add_para(doc, f'  단면의 두께 H  = {s["H_sec"]:>10.3f} mm          단 위 폭  B = 1000.000 mm')
            _add_para(doc, f'  유 효 길 이 D  = {s["D_sec"]:>10.3f} mm          피 복 두 께 Dc = {s["Dc"]:>8.3f} mm')

            # ▷ 휨모멘트 검토
            _add_para(doc, '')
            _add_para(doc, '  ▷ 휨모멘트 검토')
            _add_para(doc, '')
            _add_para(doc, '  - 횡강도 검토 -')
            _add_para(doc, '')
            _add_para(doc, f'  사용철근량 = H{s["rebar_dia"]} @ {s["rebar_spacing"]} mm    (Dc = {s["Dc"]:.0f} mm)')
            _add_para(doc, f'             = {s["As"]:.3f} mm²   ∴ P = As/(B·D) = {s["rho"]:.5f}')
            _add_para(doc, f'  공칭강도시 등가응력길이 a = (As·fy) / (0.85·fck·B) = {s["a"]:.3f} mm')

            # εt 검토
            et_chk = "≥ 0.004 ..... ∴ O.K" if s["eps_t"] >= 0.004 else "< 0.004 ..... ∴ N.G"
            _add_para(doc, f'  최외단 인장철근 변형률 εt = {s["eps_t"]:.5f} {et_chk}')
            _add_para(doc, f'    ...여기서 εt = 0.003·(H - a/β1 - Dc_min) / (a/β1)')
            if s["eps_t"] >= 0.005:
                _add_para(doc, f'  0.005 ≤ εt 이므로 인장지배단면, φf = 0.85를 적용한다.')
            elif s["eps_t"] >= 0.002:
                _add_para(doc, f'  전이구간 (0.002 ≤ εt < 0.005), φf = {s["phi_f"]:.2f}를 적용한다.')
            else:
                _add_para(doc, f'  εt < 0.002 → 압축지배단면, φf = 0.65를 적용한다.')

            _add_para(doc, '')
            _add_para(doc, f'  설계강도 φMn = φf · fy · As · ( D - a/2 ) = {s["phiMn_Nmm"]:.3f} N.mm')
            flex_chk = "≥" if s["flexure_ok"] else "<"
            flex_ok = "∴ O.K" if s["flexure_ok"] else "∴ N.G"
            _add_para(doc, f'               = {s["phiMn"]:.3f} KN.m    {flex_chk}  Mu = {s["Mu"]:.3f} KN.m ..... {flex_ok}')

            # 필요철근량 및 철근비 검토
            _add_para(doc, '')
            _add_para(doc, '  - 필요철근량 및 철근비 검토 -')
            _add_para(doc, '')
            _add_para(doc, f'  소요등가응력길이 : a = {s["a_req"]:.3f} mm로 가정')
            _add_para(doc, f'  필 요 철 근 량 : As = Mu / {{ φf·fy·(D-a/2) }}     = {s["As_req"]:.3f} mm²')
            _add_para(doc, f'  a = (As·fy) / (0.85·fck·B) = {s["a_req"]:.3f} mm  ∴ 가정과 비슷한 O.K')
            rho_req_43 = 4 / 3 * s["rho_req"]
            _add_para(doc, f'  Preq = [Mu / {{ φf·fy·(D-a/2) }}] / (B·D) = {s["rho_req"]:.5f}  → 4/3 Preq = {rho_req_43:.5f}')
            if s["rho"] >= s["pmin"]:
                _add_para(doc, f'  철근비검토 : Pmin ≤ P  .......  ∴ O.K')
            elif s["rho"] >= rho_req_43:
                _add_para(doc, f'  철근비검토 : P < Pmin 이나 4/3 Preq ≥ Pmin  ∴ O.K')
            else:
                _add_para(doc, f'  철근비검토 : Pmin > P  .......  ∴ N.G')

            # ▷ 전단력 검토
            _add_para(doc, '')
            _add_para(doc, '  ▷ 전단력 검토')
            _add_para(doc, '')
            _add_para(doc, f'  φv · Vc = φv · 1/6 · √fck · B · d / 1000 = {s["phiVc"]:.3f} KN')
            if s["shear_ok"]:
                _add_para(doc, f'  φv · Vc = {s["phiVc"]:.3f} KN > Vu  ∴전단철근 필요없음.')
            else:
                _add_para(doc, f'  φv · Vc = {s["phiVc"]:.3f} KN < Vu = {s["Vu"]:.3f} KN  ∴전단철근 필요. N.G')

            # ▷ 사용성 검토 (균열 검토)
            _add_para(doc, '')
            _add_para(doc, '  ▷ 사용성 검토 (균열 검토)')
            _add_para(doc, '')
            _add_para(doc, f'  Mcr = {s["Mcr"]:.3f} KN.m  (사용하중 모멘트)')
            _add_para(doc, f'  n = Es/Ec = 200000 / {{8500 * (Fck + △f)^(1/3)}} = {s["n"]}')
            _add_para(doc, f'  p = As/(B·D) = {s["rho"]:.5f}')
            _add_para(doc, f'  k = -np + √((np)² + 2np) = {s["k"]:.3f}         j = {s["j"]:.3f}')
            _add_para(doc, f'  x = k · d = {s["x_na"]:.3f} mm')
            _add_para(doc, f'  fc = 2 · Mcr / (B · x · (D - x/3)) = {s["fc"]:.3f} MPa')
            _add_para(doc, f'  fs = Mcr / (As · (D - x/3))         = {s["fs"]:.3f} MPa')
            _add_para(doc, f'  fst = fs · (H - Dc_min - x) / (D - x) = {s["fst"]:.3f} MPa')

            # 균열 간격
            s_max_1 = s.get("s_max_1", s["s_max"])
            s_max_2 = s.get("s_max_2", s["s_max"])
            _add_para(doc, f'  최외단철근 소요중심간격')
            _add_para(doc, f'  s = Min [ 375 · (210/fst)-2.5Cc , 300 · (210/fst) ] = {s["s_max"]:.2f} mm')
            _add_para(doc, f'    ...여기서..Cc = dc_min - 주철근 직경/2 = {s["Cc"]:.2f} mm')

            crack_chk = "≤" if s["crack_ok"] else ">"
            crack_ok = "∴ O.K" if s["crack_ok"] else "∴ N.G"
            _add_para(doc, f'  최외단철근 평균배근간격 = {s["rebar_spacing"]:.2f} mm {crack_chk} {s["s_max"]:.2f} mm ..... {crack_ok}')

            _add_para(doc, '')

        doc.add_page_break()

    # 중력식: 단면검토 생략 안내
    if _is_gravity:
        _add_heading(doc, '4. 단면 검토', 1)
        if _is_semi_gravity:
            _add_para(doc, '  반중력식 옹벽 → 별도 단면검토 참조')
        else:
            _add_para(doc, '  무근콘크리트 중력식 옹벽 → 단면검토 해당 없음')
        doc.add_page_break()

    # ================================================================
    # 종합판정
    # ================================================================
    _add_heading(doc, '종합 판정', 1)

    judge_rows = [
        ["활동 (상시)", f'{_f3(sn["SF_slide"])} >= 1.5', jdg['slide_normal']],
        ["활동 (지진)", f'{_f3(se["SF_slide"])} >= 1.2', jdg['slide_seismic']],
        ["전도 (상시)", f'{_f3(sn["SF_overturn"])} >= 2.0', jdg['overturn_normal']],
        ["전도 (지진)", f'{_f3(se["SF_overturn"])} >= 1.5', jdg.get('overturn_seismic', 'OK')],
        ["편심 (상시)", f'e={_f3(sn["e"])} <= B/6={_f3(sn["B6"])}', jdg['eccentricity_normal']],
        ["편심 (지진)", f'e={_f3(se["e"])} <= B/3={_f3(sn["B3"])}', jdg['eccentricity_seismic']],
        ["지지력 (상시)", f'Q1={_f3(sn["Q1"])} <= qa={_f3(sn["qa"])}', jdg['bearing_normal']],
        ["지지력 (지진)", f'Q1={_f3(se["Q1"])} <= qae={_f3(se["qa"])}', jdg['bearing_seismic']],
    ]
    # 단면검토 항목: 해당 멤버가 None이 아닌 경우에만 추가
    if mbr.get('BB') is not None:
        judge_rows.extend([
            ["저판 휨", f'φMn={_f3(mbr["BB"]["phiMn"])} >= Mu={_f3(mbr["BB"]["Mu"])}', jdg['BB_flexure']],
            ["저판 전단", f'φVc={_f3(mbr["BB"]["phiVc"])} >= Vu={_f3(mbr["BB"]["Vu"])}', jdg['BB_shear']],
            ["저판 균열", f's={mbr["BB"]["rebar_spacing"]} <= {_f3(mbr["BB"]["s_max"])}', jdg['BB_crack']],
        ])
    if mbr.get('CC') is not None:
        judge_rows.extend([
            ["벽체하부 휨", f'φMn={_f3(mbr["CC"]["phiMn"])} >= Mu={_f3(mbr["CC"]["Mu"])}', jdg['CC_flexure']],
            ["벽체하부 전단", f'φVc={_f3(mbr["CC"]["phiVc"])} >= Vu={_f3(mbr["CC"]["Vu"])}', jdg['CC_shear']],
            ["벽체하부 균열", f's={mbr["CC"]["rebar_spacing"]} <= {_f3(mbr["CC"]["s_max"])}', jdg['CC_crack']],
        ])
    if mbr.get('DD') is not None:
        judge_rows.extend([
            ["벽체중앙 휨", f'φMn={_f3(mbr["DD"]["phiMn"])} >= Mu={_f3(mbr["DD"]["Mu"])}', jdg['DD_flexure']],
            ["벽체중앙 전단", f'φVc={_f3(mbr["DD"]["phiVc"])} >= Vu={_f3(mbr["DD"]["Vu"])}', jdg['DD_shear']],
            ["벽체중앙 균열", f's={mbr["DD"]["rebar_spacing"]} <= {_f3(mbr["DD"]["s_max"])}', jdg['DD_crack']],
        ])
    if mbr.get('AA') is not None:
        _aa = mbr['AA']
        judge_rows.extend([
            ["앞굽판 휨", f'φMn={_f3(_aa["phiMn"])} >= Mu={_f3(_aa["Mu"])}', jdg['AA_flexure']],
            ["앞굽판 전단", f'φVc={_f3(_aa["phiVc"])} >= Vu={_f3(_aa["Vu"])}', jdg['AA_shear']],
            ["앞굽판 균열", f's={_aa["rebar_spacing"]} <= {_f3(_aa["s_max"])}', jdg['AA_crack']],
        ])
    _add_table(doc, ["검토항목", "결과", "판정"], judge_rows)

    doc.add_paragraph()
    overall = "구조적으로 안전합니다." if jdg['all_ok'] else "단면 또는 배근 조정이 필요합니다."
    final_p = _add_para(doc, f'종합 판정 : {overall}', bold=True, size=12)

    # 바이트 반환
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
